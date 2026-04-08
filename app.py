import os
import time
import threading
import telebot
import re
import io
import base64
from openai import OpenAI
from flask import Flask
from docx import Document
from PIL import Image

# --- 1. ВЕБ-СЕРВЕР ДЛЯ RENDER ---
web_app = Flask(__name__)
@web_app.route('/')
def health_check(): return "Precision-Splitter-OCR Active", 200

def run_web():
    port = int(os.environ.get("PORT", 10000))
    web_app.run(host="0.0.0.0", port=port)

# --- 2. КОНФІГУРАЦІЯ ---
TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN")
AI_KEY = os.environ.get("OPENAI_API_KEY")
bot = telebot.TeleBot(TOKEN)
client = OpenAI(api_key=AI_KEY)

user_sessions = {}
sessions_lock = threading.Lock()

# --- 3. ГОЛОВНА ЛОГІКА: РОЗРІЗАННЯ ТА СКЛЕЮВАННЯ ---
def process_fragmented_data(chat_id):
    with sessions_lock:
        session = user_sessions.get(chat_id)
        if not session: return
        image_data_list = session['images']

    try:
        raw_accumulated_text = ""
        
        for img_bytes in image_data_list:
            img = Image.open(io.BytesIO(img_bytes))
            width, height = img.size
            
            # Налаштування: 3 частини, напуск 150 пікселів (overlap)
            overlap = 150
            part_h = height // 3
            
            coords = [
                (0, 0, width, part_h + overlap),
                (0, part_h, width, 2 * part_h + overlap),
                (0, 2 * part_h, width, height)
            ]
            
            for i, box in enumerate(coords):
                fragment = img.crop(box)
                buffered = io.BytesIO()
                fragment.save(buffered, format="JPEG")
                img_b64 = base64.b64encode(buffered.getvalue()).decode('utf-8')

                # ЕТАП 1: Дослівне зчитування фрагмента
                try:
                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{
                            "role": "user",
                            "content": [
                                {"type": "text", "text": "ПЕРЕПИШИ ВЕСЬ ТЕКСТ ДОСЛІВНО. Тільки OCR знаків. Нічого не аналізуй."},
                                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_b64}", "detail": "high"}}
                            ]
                        }],
                        max_tokens=1500,
                        temperature=0
                    )
                    raw_accumulated_text += response.choices[0].message.content + "\n"
                    
                    # ПАУЗА 2 сек, щоб не отримати помилку 429 (Rate Limit)
                    time.sleep(2)

                except Exception as e:
                    if "rate_limit_exceeded" in str(e).lower():
                        bot.send_message(chat_id, "⏳ Перевищено ліміт токенів. Чекаю 10 секунд...")
                        time.sleep(10)
                    else:
                        print(f"Помилка фрагмента: {e}")

        # ЕТАП 2: Фінальна збірка (Тільки видалення дублікатів та структура)
        final_struct = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{
                "role": "user",
                "content": f"""Ось фрагменти тексту. Деякі рядки повторюються через накладання частин — ВИДАЛИ ДУБЛІКАТИ. 
                Впорядкуй текст за групами ### і нічого не скорочуй. Тільки чистий текст із документа.

                ГРУПИ ДЛЯ СТРУКТУРИ:
                ### ПРИЗНАЧЕННЯ ТА ЗАГАЛЬНИЙ ОПИС
                ### ТЕХНІЧНІ ПАРАМЕТРИ ТА ПОКАЗНИКИ
                ### ДЕТАЛЬНИЙ ОПИС КОНСТРУКЦІЇ
                ### ПОРЯДОК РОБОТИ ТА ОБСЛУГОВУВАННЯ
                ### ВІЗУАЛЬНІ ДАНІ ТА ПРИМІТКИ

                ТЕКСТ:
                {raw_accumulated_text}"""
            }]
        )
        
        clean_text = final_struct.choices[0].message.content

        # ЕТАП 3: Створення файлу Word
        doc = Document()
        doc.add_heading("Оцифрований документ", 0)
        
        for line in clean_text.split('\n'):
            line = line.strip()
            if not line: continue
            
            if line.startswith('###'):
                doc.add_heading(line.replace('###', '').strip(), level=1)
            elif ":" in line and len(line.split(":")[0]) < 65:
                p = doc.add_paragraph(style='List Bullet')
                parts = line.split(":", 1)
                p.add_run(parts[0].strip() + ": ").bold = True
                p.add_run(parts[1].strip())
            else:
                doc.add_paragraph(line)

        file_path = f"report_{chat_id}.docx"
        doc.save(file_path)
        
        with open(file_path, "rb") as f:
            bot.send_document(chat_id, f, caption="✅ Оцифровка завершена. Обсяг збережено, дублікати видалено.")
        
        if os.path.exists(file_path): os.remove(file_path)

    except Exception as e:
        bot.send_message(chat_id, f"❌ Критична помилка: {e}")
    finally:
        with sessions_lock:
            user_sessions.pop(chat_id, None)

# --- 4. ОБРОБНИКИ ТЕЛЕГРАМ ---
@bot.message_handler(content_types=['photo'])
def handle_photos(message):
    chat_id = message.chat.id
    file_info = bot.get_file(message.photo[-1].file_id)
    downloaded_file = bot.download_file(file_info.file_path)

    with sessions_lock:
        if chat_id not in user_sessions:
            user_sessions[chat_id] = {'images': [], 'timer': None}
            bot.send_message(chat_id, "⚙️ Запущено фрагментарне зчитування (обхід лімітів та фільтрів)...")
        
        user_sessions[chat_id]['images'].append(downloaded_file)
        
        if user_sessions[chat_id]['timer']:
            user_sessions[chat_id]['timer'].cancel()
        
        # Таймер 12 секунд для збору альбому
        t = threading.Timer(12.0, process_fragmented_data, args=[chat_id])
        user_sessions[chat_id]['timer'] = t
        t.start()

if __name__ == "__main__":
    threading.Thread(target=run_web, daemon=True).start()
    bot.remove_webhook()
    time.sleep(1)
    print("Бот запущений (Precision Mode + RateLimit Protection)")
    bot.infinity_polling(timeout=90)
